from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


HEADING_RE = re.compile(r'^(#{1,6})\s+(.*)$')
UNORDERED_LIST_RE = re.compile(r'^\s*[-*]\s+(.*)$')
ORDERED_LIST_RE = re.compile(r'^\s*\d+\.\s+(.*)$')
IMAGE_RE = re.compile(r'^!\[([^\]]*)\]\(([^)]+)\)$')


def get_style_name(document: Document, preferred: str, fallback: str = 'Normal') -> str:
    try:
        document.styles[preferred]
        return preferred
    except KeyError:
        return fallback


def split_table_row(line: str) -> list[str]:
    return [part.strip() for part in line.strip().strip('|').split('|')]


def is_table_separator(line: str) -> bool:
    parts = split_table_row(line)
    return bool(parts) and all(re.fullmatch(r':?-{3,}:?', part) for part in parts)


def add_inline_runs(paragraph, text: str) -> None:
    code_segments = re.split(r'(`[^`]+`)', text)
    for segment in code_segments:
        if not segment:
            continue
        if segment.startswith('`') and segment.endswith('`'):
            run = paragraph.add_run(segment[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            continue

        bold_segments = re.split(r'(\*\*[^*]+\*\*)', segment)
        for bold_segment in bold_segments:
            if not bold_segment:
                continue
            is_bold = bold_segment.startswith('**') and bold_segment.endswith('**')
            run = paragraph.add_run(bold_segment[2:-2] if is_bold else bold_segment)
            run.bold = is_bold


def add_text_paragraph(document: Document, text: str, style_name: str = 'Normal'):
    paragraph = document.add_paragraph(style=get_style_name(document, style_name))
    add_inline_runs(paragraph, text)
    return paragraph


def add_code_block(document: Document, code_lines: list[str]) -> None:
    paragraph = document.add_paragraph(style=get_style_name(document, 'No Spacing'))
    for index, code_line in enumerate(code_lines):
        run = paragraph.add_run(code_line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        if index < len(code_lines) - 1:
            run.add_break()


def add_table(document: Document, table_lines: list[str]) -> None:
    rows = [split_table_row(line) for line in table_lines]
    if len(rows) >= 2 and is_table_separator(table_lines[1]):
        headers = rows[0]
        data_rows = rows[2:]
    else:
        headers = rows[0]
        data_rows = rows[1:]

    column_count = max(len(headers), *(len(row) for row in data_rows), 1)
    table = document.add_table(rows=1 + len(data_rows), cols=column_count)
    table.style = 'Table Grid'

    for column_index, cell_text in enumerate(headers):
        paragraph = table.cell(0, column_index).paragraphs[0]
        add_inline_runs(paragraph, cell_text)
        for run in paragraph.runs:
            run.bold = True

    for row_index, row in enumerate(data_rows, start=1):
        for column_index, cell_text in enumerate(row):
            paragraph = table.cell(row_index, column_index).paragraphs[0]
            add_inline_runs(paragraph, cell_text)


def convert_markdown(markdown_path: Path, output_path: Path) -> None:
    document = Document()
    document.core_properties.title = markdown_path.stem

    normal_style = document.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)

    lines = markdown_path.read_text(encoding='utf-8').splitlines()
    index = 0
    first_heading_written = False

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            if not first_heading_written and level == 1:
                paragraph = document.add_heading(text, 0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                first_heading_written = True
            else:
                document.add_heading(text, level=min(level, 4))
            index += 1
            continue

        if stripped.startswith('```'):
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith('```'):
                code_lines.append(lines[index].rstrip())
                index += 1
            add_code_block(document, code_lines)
            if index < len(lines):
                index += 1
            continue

        image_match = IMAGE_RE.match(stripped)
        if image_match:
            image_path = (markdown_path.parent / image_match.group(2)).resolve()
            if image_path.exists():
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run().add_picture(str(image_path), width=Inches(6.0))
                if image_match.group(1).strip():
                    caption = document.add_paragraph(style=get_style_name(document, 'Caption'))
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_inline_runs(caption, image_match.group(1).strip())
            else:
                add_text_paragraph(document, f'[Missing image: {image_match.group(2)}]')
            index += 1
            continue

        if stripped.startswith('|'):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith('|'):
                table_lines.append(lines[index])
                index += 1
            add_table(document, table_lines)
            continue

        unordered_match = UNORDERED_LIST_RE.match(stripped)
        if unordered_match:
            while index < len(lines):
                current = lines[index].strip()
                current_match = UNORDERED_LIST_RE.match(current)
                if not current_match:
                    break
                add_text_paragraph(document, current_match.group(1).strip(), 'List Bullet')
                index += 1
            continue

        ordered_match = ORDERED_LIST_RE.match(stripped)
        if ordered_match:
            while index < len(lines):
                current = lines[index].strip()
                current_match = ORDERED_LIST_RE.match(current)
                if not current_match:
                    break
                add_text_paragraph(document, current_match.group(1).strip(), 'List Number')
                index += 1
            continue

        paragraph_parts: list[str] = []
        while index < len(lines):
            current = lines[index].strip()
            if not current:
                break
            if (
                HEADING_RE.match(current)
                or current.startswith('```')
                or current.startswith('|')
                or IMAGE_RE.match(current)
                or UNORDERED_LIST_RE.match(current)
                or ORDERED_LIST_RE.match(current)
            ):
                break
            paragraph_parts.append(current)
            index += 1

        add_text_paragraph(document, ' '.join(paragraph_parts))

    document.save(output_path)


def main() -> int:
    if len(sys.argv) != 3:
        print('Usage: python generate_docx_report.py <input.md> <output.docx>')
        return 1

    input_path = Path(sys.argv[1]).resolve()
    output_path = Path(sys.argv[2]).resolve()
    convert_markdown(input_path, output_path)
    print(output_path)
    return 0


if __name__ == '__main__':
    raise SystemExit(main())